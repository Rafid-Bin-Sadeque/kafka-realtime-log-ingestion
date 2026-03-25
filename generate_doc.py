from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Helpers ───────────────────────────────────────────────────────────────────
def h1(text):
    return doc.add_heading(text, level=1)

def h2(text):
    return doc.add_heading(text, level=2)

def h3(text):
    return doc.add_heading(text, level=3)

def body(text):
    return doc.add_paragraph(text)

def bullet(text):
    return doc.add_paragraph(text, style="List Bullet")

def numbered(text):
    return doc.add_paragraph(text, style="List Number")

def code_block(text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2E, 0x86, 0x0E)
    p.paragraph_format.left_indent = Inches(0.4)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F4F4F4")
    p._p.get_or_add_pPr().append(shading)
    return p

def note(text):
    p = doc.add_paragraph()
    run = p.add_run("NOTE: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0xCC, 0x55, 0x00)
    p.add_run(text)
    return p

def tip(text):
    p = doc.add_paragraph()
    run = p.add_run("TIP: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)
    p.add_run(text)
    return p

# ── Title Page ────────────────────────────────────────────────────────────────
title = doc.add_heading("Kafka from Scratch", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph("A plain-English guide to Docker, Kafka, and everything in your project")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True
doc.add_page_break()

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 1 — WHAT IS DOCKER
# ═════════════════════════════════════════════════════════════════════════════
h1("Chapter 1 — What is Docker?")

h2("The problem Docker solves")
body(
    "Imagine you write a Python app on your laptop that depends on a specific version of Kafka, "
    "Postgres, and a handful of Python packages. It works perfectly on your machine. "
    "You send it to a colleague and it immediately crashes — they have different versions installed, "
    "a different OS, or a conflicting library. This is the classic 'works on my machine' problem."
)
body(
    "Docker solves this by packaging the application and everything it needs — the runtime, "
    "libraries, config, environment variables — into a single portable unit called a container. "
    "You ship the container, not just the code. It runs identically everywhere."
)

h2("Containers vs Virtual Machines — what is the difference?")
body(
    "Both containers and Virtual Machines (VMs) isolate software from the host system, "
    "but they do it in very different ways."
)
body(
    "A Virtual Machine runs a full copy of an operating system on top of your OS, "
    "managed by a hypervisor. It is like running a second computer inside your computer. "
    "This is powerful but heavy: a VM image is typically several gigabytes, "
    "and booting one takes minutes."
)
body(
    "A container does NOT run its own OS. Instead it shares the host machine's OS kernel "
    "but isolates the process inside its own filesystem, network, and process space. "
    "A container image is typically tens to hundreds of megabytes, and it starts in seconds."
)
bullet("VM: full OS inside every machine — slow to start, heavy on disk and RAM")
bullet("Container: shares the host kernel, isolates just the app — fast, lightweight")
bullet("Both give you isolation; containers give it with much less overhead")

h2("The core Docker vocabulary — learn these five terms")
body("Everything in Docker comes back to these concepts:")
body("")

p = doc.add_paragraph()
r = p.add_run("Image")
r.bold = True
p.add_run(
    " — A read-only template that describes what a container should contain: "
    "the OS base layer, the app binaries, config files, environment variables. "
    "Think of it like a class in object-oriented programming — a blueprint, not a running thing. "
    "Images are built once and can be used to create many containers. "
    "Images are stored in a registry (Docker Hub is the public default)."
)

p = doc.add_paragraph()
r = p.add_run("Container")
r.bold = True
p.add_run(
    " — A running instance of an image. Like an object created from a class. "
    "You can create many containers from the same image. "
    "Each container has its own isolated filesystem, network interface, and process space. "
    "When a container stops, everything written to its filesystem is lost "
    "unless you explicitly saved it to a volume."
)

p = doc.add_paragraph()
r = p.add_run("Volume")
r.bold = True
p.add_run(
    " — A directory on your real (host) machine that is mounted into a container. "
    "Data written to a volume persists after the container stops or is deleted. "
    "This is how you persist Kafka's message log across container restarts. "
    "Without a volume, all data inside a container is ephemeral."
)

p = doc.add_paragraph()
r = p.add_run("Network")
r.bold = True
p.add_run(
    " — Docker creates a private virtual network for your containers. "
    "Containers on the same network can reach each other by their container name "
    "(e.g. the 'kafka' container is reachable at hostname 'kafka' from other containers). "
    "This is how Kafka UI finds the Kafka broker without needing to know an IP address."
)

p = doc.add_paragraph()
r = p.add_run("Port mapping")
r.bold = True
p.add_run(
    " — A container's ports are private by default — your laptop cannot reach them. "
    "A port mapping bridges a port inside the container to a port on your laptop. "
    "The syntax is HOST_PORT:CONTAINER_PORT. For example '9092:9092' means "
    "traffic hitting localhost:9092 on your machine is forwarded into the container's port 9092."
)

h2("How Docker Hub fits in")
body(
    "Docker Hub (hub.docker.com) is the default public registry for Docker images. "
    "When you use an image name like 'confluentinc/cp-kafka:7.4.0', Docker breaks it down as:"
)
bullet("confluentinc — the organisation/user account on Docker Hub")
bullet("cp-kafka — the image name")
bullet("7.4.0 — the tag (version). If you omit the tag, Docker uses 'latest'")
body(
    "When you run a container for the first time, Docker automatically pulls the image "
    "from Docker Hub if it is not already cached on your machine."
)

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 2 — STANDALONE DOCKER COMMANDS
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Chapter 2 — Docker Commands: The Building Blocks")

body(
    "Before we get to Docker Compose (which manages multiple containers at once), "
    "it is important to understand the individual 'docker' commands that Compose uses under the hood. "
    "Every 'docker compose' action is ultimately calling these commands for you."
)

h2("Working with images")

h3("docker pull — download an image")
body(
    "Downloads an image from Docker Hub (or another registry) to your local machine. "
    "You do not need to run this manually — 'docker run' pulls automatically if the image is missing. "
    "But explicit pulls are useful when you want to pre-download or update an image."
)
code_block(
    "docker pull confluentinc/cp-kafka:7.4.0\n"
    "# Downloads the Confluent Kafka image, version 7.4.0\n\n"
    "docker pull confluentinc/cp-kafka\n"
    "# Downloads 'latest' tag if no version is specified"
)

h3("docker images — list downloaded images")
body("Shows all images stored locally on your machine.")
code_block(
    "docker images\n\n"
    "# Output looks like:\n"
    "# REPOSITORY                 TAG       IMAGE ID       SIZE\n"
    "# confluentinc/cp-kafka      7.4.0     a1b2c3d4e5f6   800MB\n"
    "# confluentinc/cp-zookeeper  latest    f6e5d4c3b2a1   600MB"
)

h3("docker rmi — remove an image")
body("Deletes an image from your local disk. Useful to free up space.")
code_block(
    "docker rmi confluentinc/cp-kafka:7.4.0\n\n"
    "# You cannot remove an image if a container (even a stopped one) is using it.\n"
    "# Remove the container first, then the image."
)

h2("Working with containers")

h3("docker run — create and start a container")
body(
    "This is the most important command. It creates a container from an image and starts it. "
    "There are many flags — here are the ones you will use most:"
)
code_block(
    "docker run hello-world\n"
    "# Simplest possible run: pulls 'hello-world' image and runs it once, then exits"
)
code_block(
    "docker run -d --name my-kafka -p 9092:9092 confluentinc/cp-kafka:7.4.0\n"
    "#          ^  ^               ^             ^\n"
    "#          |  |               |             image to use\n"
    "#          |  |               map host port 9092 to container port 9092\n"
    "#          |  give the container a name ('my-kafka') instead of a random one\n"
    "#          run in background (detached mode)"
)
code_block(
    "docker run -it ubuntu bash\n"
    "# -i = interactive (keep stdin open)\n"
    "# -t = allocate a terminal (TTY)\n"
    "# Together -it gives you an interactive shell inside the container"
)
code_block(
    "docker run --rm ubuntu echo 'hello'\n"
    "# --rm = automatically delete the container when it exits\n"
    "# Useful for one-off commands where you do not want leftover containers"
)
code_block(
    "docker run -e MY_VAR=hello ubuntu\n"
    "# -e = set an environment variable inside the container\n"
    "# Same as putting it under 'environment:' in docker-compose.yml"
)
code_block(
    "docker run -v /my/local/folder:/data ubuntu\n"
    "# -v = mount a volume\n"
    "# /my/local/folder on your machine is visible as /data inside the container"
)

h3("docker ps — list running containers")
body("Shows containers that are currently running.")
code_block(
    "docker ps\n\n"
    "# Output:\n"
    "# CONTAINER ID   IMAGE                  COMMAND    STATUS    PORTS                    NAMES\n"
    "# a1b2c3d4e5f6   confluentinc/cp-kafka  ...        Up        0.0.0.0:9092->9092/tcp   kafka-kafka-1"
)
code_block(
    "docker ps -a\n"
    "# -a = show ALL containers, including stopped ones\n"
    "# Without -a you only see running containers"
)

h3("docker stop — gracefully stop a container")
body(
    "Sends a SIGTERM signal to the main process inside the container, giving it time to "
    "shut down cleanly. After a timeout (default 10 seconds) it sends SIGKILL if still running."
)
code_block(
    "docker stop kafka-kafka-1\n"
    "# Stop by container name\n\n"
    "docker stop a1b2c3d4e5f6\n"
    "# Stop by container ID (can use just the first few characters)"
)

h3("docker start — restart a stopped container")
body("Starts a container that was previously stopped. The container keeps its original config.")
code_block("docker start kafka-kafka-1")

h3("docker rm — remove a container")
body(
    "Deletes a stopped container. Does NOT delete the image. "
    "A container must be stopped before you can remove it."
)
code_block(
    "docker rm kafka-kafka-1\n\n"
    "docker rm -f kafka-kafka-1\n"
    "# -f = force remove even if the container is still running (stop + remove in one step)"
)

h3("docker exec — run a command inside a running container")
body(
    "This is how you 'SSH into' a container or run a one-off command inside it. "
    "The container must already be running."
)
code_block(
    "docker exec -it kafka-kafka-1 bash\n"
    "# Open an interactive bash shell inside the kafka container\n"
    "# -i = interactive, -t = terminal\n"
    "# Type 'exit' to leave without stopping the container"
)
code_block(
    "docker exec kafka-kafka-1 ls /var/lib/kafka/data\n"
    "# Run a single command and get the output (no interactive shell needed)"
)
code_block(
    "docker exec -it kafka-kafka-1 kafka-topics --bootstrap-server localhost:9092 --list\n"
    "# Run the kafka-topics CLI tool that lives inside the container"
)
note(
    "This is why Kafka admin commands require 'docker exec' — the kafka-topics CLI tool "
    "is only installed inside the container, not on your laptop."
)

h3("docker logs — view a container's output")
body("Shows the stdout/stderr of a running or stopped container.")
code_block(
    "docker logs kafka-kafka-1\n"
    "# Print all logs so far\n\n"
    "docker logs -f kafka-kafka-1\n"
    "# -f = follow (live tail, like 'tail -f')\n\n"
    "docker logs --tail 50 kafka-kafka-1\n"
    "# Only show the last 50 lines"
)

h3("docker inspect — get detailed info about a container or image")
body(
    "Returns a large JSON blob with every detail about a container: "
    "its network config, mounted volumes, environment variables, restart policy, etc."
)
code_block(
    "docker inspect kafka-kafka-1\n\n"
    "# Useful for debugging — e.g. to confirm which port is actually mapped:\n"
    "docker inspect kafka-kafka-1 | grep -A 5 'PortBindings'"
)

h2("Working with volumes")

h3("docker volume ls — list volumes")
code_block("docker volume ls")

h3("docker volume inspect — see where a volume lives on disk")
code_block(
    "docker volume inspect kafka_kafka_data\n"
    "# Shows the actual path on your Mac/Linux machine where the data is stored"
)

h3("docker volume rm — delete a volume")
code_block(
    "docker volume rm kafka_kafka_data\n"
    "# WARNING: this permanently deletes all data in the volume\n"
    "# The volume must not be in use by any container"
)

h2("Working with networks")

h3("docker network ls — list networks")
code_block("docker network ls")

h3("docker network inspect — see which containers are on a network")
code_block(
    "docker network inspect kafka_default\n"
    "# Shows all containers connected to this network and their internal IPs"
)

h2("Cleaning up everything at once")
body("Docker accumulates stopped containers, unused images, and orphaned volumes over time. Use these to clean up:")
code_block(
    "docker system prune\n"
    "# Remove all stopped containers, unused networks, dangling images\n\n"
    "docker system prune -a\n"
    "# Also remove unused images (not just dangling ones) — frees a lot of disk\n\n"
    "docker system prune --volumes\n"
    "# Also remove unused volumes — BE CAREFUL, this deletes data"
)

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 3 — DOCKER COMPOSE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Chapter 3 — Docker Compose: Running Multiple Containers Together")

h2("Why Compose exists")
body(
    "Running a single container with 'docker run' is fine. "
    "But running Kafka requires three containers (Zookeeper, Kafka, Kafka UI) that all need to "
    "know about each other, start in the right order, and share a network. "
    "Doing this with individual 'docker run' commands would mean remembering a long series of "
    "flags and running them in the right sequence every time."
)
body(
    "Docker Compose lets you describe your entire multi-container setup in a single YAML file "
    "(docker-compose.yml) and then manage it all with simple commands."
)

h2("The docker-compose.yml file — anatomy")
body(
    "Every docker-compose.yml has the same structure. Here is what each section means:"
)
code_block(
    "services:              # Define each container you want to run\n"
    "  my-service:          # The service name — also its hostname on the Docker network\n"
    "    image: ...         # Which Docker image to use\n"
    "    build: .           # OR: build from a local Dockerfile instead of pulling\n"
    "    ports:             # Port mappings: HOST:CONTAINER\n"
    "      - '8080:80'\n"
    "    environment:       # Environment variables to set inside the container\n"
    "      KEY: value\n"
    "    volumes:           # Mount host directories or named volumes\n"
    "      - ./local:/container/path\n"
    "      - named_vol:/data\n"
    "    depends_on:        # Start this service only after these others are started\n"
    "      - other-service\n"
    "    networks:          # Which networks to join (optional — default network is created)\n"
    "      - my-network\n\n"
    "volumes:               # Declare named volumes here\n"
    "  named_vol:\n\n"
    "networks:              # Declare custom networks here (optional)\n"
    "  my-network:"
)

h2("Docker Compose commands")

h3("docker compose up — start all services")
code_block(
    "docker compose up\n"
    "# Start all services, print logs to terminal\n"
    "# Ctrl+C stops everything\n\n"
    "docker compose up -d\n"
    "# -d = detached: run in background, return your terminal\n\n"
    "docker compose up kafka\n"
    "# Start only the 'kafka' service (and its dependencies)"
)
body("What happens when you run 'docker compose up -d':")
numbered("Docker reads docker-compose.yml in the current directory.")
numbered("For each service, it checks if the image is cached locally. If not, it pulls from Docker Hub.")
numbered("It creates a private Docker network named '<folder>_default' (e.g. 'kafka_default').")
numbered("It starts containers in dependency order: zookeeper first (no dependencies), then kafka (depends_on zookeeper), then kafka-ui (depends_on kafka).")
numbered("Port mappings are applied so localhost:9092, localhost:2181, localhost:8080 reach the right containers.")
numbered("Named volumes are created on first run and mounted into the containers.")

h3("docker compose down — stop and remove everything")
code_block(
    "docker compose down\n"
    "# Stop and remove all containers AND the network\n"
    "# Images and volumes are NOT deleted — data in named volumes is preserved\n\n"
    "docker compose down -v\n"
    "# Also delete named volumes — ALL Kafka data is wiped\n"
    "# Use this when you want a completely fresh start\n\n"
    "docker compose down --rmi all\n"
    "# Also delete the images — forces re-download next time"
)

h3("docker compose start / stop — without removing containers")
code_block(
    "docker compose stop\n"
    "# Stop all containers but keep them (like docker stop)\n"
    "# Data is preserved, containers still exist\n\n"
    "docker compose start\n"
    "# Restart previously stopped containers"
)
note(
    "'docker compose stop' + 'docker compose start' is faster than 'down' + 'up' "
    "because it skips pulling images and recreating the network."
)

h3("docker compose restart — restart specific services")
code_block(
    "docker compose restart kafka\n"
    "# Restart just the Kafka broker without touching zookeeper or kafka-ui\n\n"
    "docker compose restart\n"
    "# Restart all services"
)

h3("docker compose ps — see status of services")
code_block(
    "docker compose ps\n\n"
    "# Output:\n"
    "# NAME             IMAGE                   STATUS    PORTS\n"
    "# kafka-kafka-1    confluentinc/cp-kafka   Up        0.0.0.0:9092->9092/tcp\n"
    "# kafka-zookeeper  confluentinc/cp-zoo...  Up        0.0.0.0:2181->2181/tcp\n"
    "# kafka-kafka-ui   provectuslabs/kafka-ui  Up        0.0.0.0:8080->8080/tcp"
)

h3("docker compose logs — view service output")
code_block(
    "docker compose logs\n"
    "# Logs from all services\n\n"
    "docker compose logs kafka\n"
    "# Logs from just the kafka service\n\n"
    "docker compose logs -f kafka\n"
    "# Follow (live tail) kafka logs\n\n"
    "docker compose logs --tail 100 kafka\n"
    "# Last 100 lines from kafka"
)

h3("docker compose exec — run a command inside a running service")
body(
    "Same as 'docker exec' but you use the service name from docker-compose.yml "
    "instead of the full container name."
)
code_block(
    "docker compose exec kafka bash\n"
    "# Open a shell in the kafka service container\n\n"
    "docker compose exec kafka kafka-topics --bootstrap-server localhost:9092 --list\n"
    "# Run a command inside the kafka container"
)

h3("docker compose build — build images from a Dockerfile")
body(
    "This project uses pre-built images from Docker Hub, so 'build' is not needed here. "
    "But when you write your own app with a Dockerfile, you use this:"
)
code_block(
    "docker compose build\n"
    "# Build all services that have a 'build:' key in docker-compose.yml\n\n"
    "docker compose build my-app\n"
    "# Build only a specific service\n\n"
    "docker compose up --build\n"
    "# Build AND start — rebuilds images before starting (useful during development)"
)

h3("docker compose pull — update images")
code_block(
    "docker compose pull\n"
    "# Pull the latest version of all images defined in docker-compose.yml\n"
    "# Does not restart containers — run 'docker compose up -d' after to apply"
)

h2("depends_on — controlling startup order")
body(
    "depends_on tells Compose to start one service before another. "
    "However, it only waits for the container to START, not for the app inside it to be READY. "
    "Kafka may start its container before Zookeeper is actually accepting connections. "
    "In production you would use healthchecks to solve this. For learning, "
    "Kafka has a built-in retry loop that keeps trying to connect to Zookeeper."
)
code_block(
    "kafka:\n"
    "  depends_on:\n"
    "    - zookeeper    # Start zookeeper container first, then start kafka"
)

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 4 — THIS PROJECT'S DOCKER SETUP
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Chapter 4 — This Project's Docker Setup")

h2("The docker-compose.yml walkthrough")
body("Here is your docker-compose.yml broken down line by line:")

h3("Zookeeper service")
code_block(
    "zookeeper:\n"
    "  image: confluentinc/cp-zookeeper:latest\n"
    "  ports:\n"
    "    - '2181:2181'          # expose Zookeeper's client port to localhost\n"
    "  environment:\n"
    "    ZOOKEEPER_CLIENT_PORT: 2181   # the port Zookeeper listens on for clients\n"
    "    ZOOKEEPER_TICK_TIME: 2000     # heartbeat interval in milliseconds\n"
    "                                  # used to calculate session/leader timeouts"
)

h3("Kafka broker service")
code_block(
    "kafka:\n"
    "  image: confluentinc/cp-kafka:7.4.0\n"
    "  ports:\n"
    "    - '9092:9092'          # expose Kafka to localhost so your Python code can connect\n"
    "  environment:\n"
    "    KAFKA_BROKER_ID: 1                        # unique ID in the cluster (only 1 broker here)\n"
    "    KAFKA_ZOOKEEPER_CONNECT: zookeeper:2181   # reach Zookeeper by service name on Docker network\n"
    "    KAFKA_ADVERTISED_LISTENERS: PLAINTEXT://localhost:9092\n"
    "    # ^ what Kafka tells clients to connect to (your laptop reaches it via localhost)\n"
    "    KAFKA_LISTENERS: PLAINTEXT://0.0.0.0:9092\n"
    "    # ^ what address Kafka actually binds to inside the container (all interfaces)\n"
    "    KAFKA_AUTO_CREATE_TOPICS_ENABLE: 'true'   # create a topic on first produce\n"
    "    KAFKA_OFFSETS_TOPIC_REPLICATION_FACTOR: 1 # 1 broker = replication factor must be 1\n"
    "  depends_on:\n"
    "    - zookeeper"
)

h3("Kafka UI service")
code_block(
    "kafka-ui:\n"
    "  image: provectuslabs/kafka-ui:latest\n"
    "  ports:\n"
    "    - '8080:8080'          # web dashboard at http://localhost:8080\n"
    "  environment:\n"
    "    - KAFKA_CLUSTERS_0_NAME=local\n"
    "    - KAFKA_CLUSTERS_0_BOOTSTRAPSERVERS=kafka:9092\n"
    "    # ^ connects to Kafka by service name (inside Docker network)\n"
    "  depends_on:\n"
    "    - kafka"
)

h2("LISTENERS vs ADVERTISED_LISTENERS — the most confusing Kafka setting")
body(
    "This trips up almost everyone. Here is a clear explanation:"
)
bullet(
    "KAFKA_LISTENERS (0.0.0.0:9092) — the socket address Kafka binds to INSIDE the container. "
    "0.0.0.0 means 'listen on all network interfaces'. This is the actual server socket."
)
bullet(
    "KAFKA_ADVERTISED_LISTENERS (localhost:9092) — what Kafka puts in its metadata response. "
    "When a producer connects and asks 'where is the leader for partition 0?', "
    "Kafka responds with this address. Your Python code then uses this address to "
    "connect directly to the partition leader."
)
body(
    "Why two settings? Because the address the broker listens on inside the container "
    "(0.0.0.0:9092) is different from the address clients should use to reach it from outside "
    "(localhost:9092 via the port mapping). If you only had one setting, "
    "clients from outside would be told to connect to 0.0.0.0:9092 which is meaningless outside the container."
)

h2("How the network flows end to end")
code_block(
    "Your Python producer.py\n"
    "  |  connects to localhost:9092  (your laptop)\n"
    "  |\n"
    "  v\n"
    "Docker port mapping 9092:9092\n"
    "  |  forwards to container port 9092\n"
    "  |\n"
    "  v\n"
    "Kafka broker  (inside Docker network 'kafka_default')\n"
    "  |  on startup, registered itself with Zookeeper at zookeeper:2181\n"
    "  |  stores your message in /var/lib/kafka/data/orders-0/\n"
    "  |\n"
    "  v\n"
    "Your Python consumer.py\n"
    "  |  also connects to localhost:9092\n"
    "  |  polls for new messages\n"
    "  v\n"
    "Kafka UI (http://localhost:8080)\n"
    "  connects to kafka:9092 (inside Docker network, no port mapping needed)"
)

h2("Recommended start workflow")
code_block(
    "# 1. Start everything\n"
    "docker compose up -d\n\n"
    "# 2. Verify all three services are running\n"
    "docker compose ps\n\n"
    "# 3. Watch Kafka finish starting up (look for 'started' in logs)\n"
    "docker compose logs -f kafka\n"
    "# Ctrl+C once you see it is ready\n\n"
    "# 4. Open Kafka UI in browser\n"
    "# http://localhost:8080\n\n"
    "# 5. Run producer, then consumer in separate terminals\n"
    "python3 01_basics/producer.py\n"
    "python3 01_basics/consumer.py\n\n"
    "# 6. When done\n"
    "docker compose down"
)

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 5 — WHAT IS KAFKA
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Chapter 5 — What is Kafka?")

h2("One-sentence definition")
body(
    "Kafka is a distributed commit log: producers append messages to it, "
    "consumers read from it at their own pace, and the log is retained on disk "
    "for a configurable amount of time regardless of whether anyone has read it."
)

h2("Why not just use a database or a queue?")
bullet(
    "A traditional message queue (RabbitMQ, SQS) deletes a message once a consumer acknowledges it. "
    "Kafka keeps it. Multiple independent apps can all read the same event."
)
bullet(
    "A database is optimised for random reads and writes. "
    "Kafka is optimised for sequential appends and sequential reads — "
    "which is orders of magnitude faster for high-volume streaming."
)
bullet(
    "Kafka lets multiple independent consumer groups each read the same data "
    "at different speeds without affecting each other."
)

h2("The mental model")
body(
    "Think of Kafka as a river of events. "
    "Producers throw messages into the river. "
    "Consumers paddle down the river reading messages. "
    "Different consumers can be at different points in the river at the same time. "
    "The river keeps flowing whether or not anyone is reading."
)

h2("Core components")
bullet("Broker — the Kafka server process. Stores the log on disk and serves reads/writes.")
bullet("Topic — a named category of messages (e.g. 'orders'). Like a table in a DB, but append-only.")
bullet("Partition — a topic is split into N partitions. Each partition is an independent ordered log.")
bullet("Producer — a client that writes messages to a topic.")
bullet("Consumer — a client that reads messages from a topic.")
bullet("Consumer Group — a set of consumers that cooperate to read a topic. Each partition is owned by exactly one consumer in the group at any time.")
bullet("Offset — the position of a message inside a partition. Starts at 0, increments by 1. Permanent and immutable.")

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 6 — WHY THREE SERVICES
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Chapter 6 — Why Three Services? (Zookeeper, Kafka, Kafka UI)")

h2("Service 1: Zookeeper")
body(
    "Zookeeper is a distributed coordination service. In the context of Kafka it acts as the "
    "cluster's brain — it keeps track of:"
)
bullet("Which brokers are alive")
bullet("Who the leader of each partition is")
bullet("Access control lists")
bullet("Topic configuration metadata")
body(
    "Without Zookeeper (in older Kafka versions), brokers have no way to agree on who is in charge "
    "of what. If the leader of a partition dies, Zookeeper coordinates electing a new one."
)
note(
    "Kafka 3.x introduced KRaft mode which removes the Zookeeper dependency — Kafka manages its "
    "own metadata internally. This project uses the older Zookeeper-based setup because it is "
    "simpler to reason about when learning."
)

h2("Service 2: Kafka Broker")
body("This is the actual Kafka server. It does the real work:")
bullet("Receives messages from producers and writes them to disk")
bullet("Serves messages to consumers")
bullet("Maintains partition leadership")
bullet("Enforces retention policies (how long to keep messages)")
body(
    "In production you run multiple brokers for fault tolerance. In this project you have one "
    "broker (KAFKA_BROKER_ID: 1) which is fine for learning."
)

h2("Service 3: Kafka UI")
body(
    "Kafka UI (provectuslabs/kafka-ui) is a web dashboard. It has no role in the data pipeline — "
    "it is purely an observability tool. It connects to your Kafka broker and lets you:"
)
bullet("Browse topics and their partitions")
bullet("Read individual messages visually")
bullet("Monitor consumer group lag")
bullet("Create or delete topics")
body("Visit it at http://localhost:8080 while the stack is running.")

h2("How the three services talk to each other")
body("Inside Docker's private network:")
code_block(
    "kafka-ui  --->  kafka:9092       (reads topic/message metadata)\n"
    "kafka     --->  zookeeper:2181   (registers itself, reads cluster state)\n"
    "your app  --->  localhost:9092   (mapped to kafka:9092 inside Docker)"
)

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 7 — WHERE THE LOG LIVES
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Chapter 7 — Where Does the Producer's Data Actually Live?")

h2("The Kafka log on disk")
body(
    "Every message a producer sends is written to disk inside the Kafka broker container. "
    "The default storage path inside the container is:"
)
code_block("/var/lib/kafka/data/")
body("Inside that directory Kafka creates a folder per topic-partition:")
code_block(
    "/var/lib/kafka/data/orders-0/   <- topic 'orders', partition 0\n"
    "/var/lib/kafka/data/orders-1/   <- topic 'orders', partition 1"
)
body("Inside each folder you will find segment files:")
code_block(
    "00000000000000000000.log        <- the actual message bytes\n"
    "00000000000000000000.index      <- maps offsets to byte positions\n"
    "00000000000000000000.timeindex  <- maps timestamps to offsets"
)

h2("Segment files explained")
bullet(".log file: raw binary file. Each entry = message metadata + key bytes + value bytes. Appended sequentially. Never modified, only appended.")
bullet(".index file: sparse index. Allows Kafka to jump to a specific offset without scanning the whole .log file.")
bullet(".timeindex file: lets Kafka find messages by timestamp (used for time-based retention and offset lookup by time).")

h2("Is the data safe across container restarts?")
body(
    "By default in this project the data lives inside the container's ephemeral filesystem. "
    "If you run 'docker compose down', the container is deleted and the data is gone. "
    "To persist data add a volume mount to docker-compose.yml:"
)
code_block(
    "kafka:\n"
    "  volumes:\n"
    "    - kafka_data:/var/lib/kafka/data\n\n"
    "volumes:\n"
    "  kafka_data:"
)
body(
    "With this in place, 'docker compose down' preserves data. "
    "Only 'docker compose down -v' deletes the volume."
)

h2("How the consumer reads it")
body(
    "The consumer does NOT read from disk directly — it connects to the broker over the network "
    "and asks for messages starting from a given offset. The broker looks up the byte position "
    "in the .index file, seeks to that position in the .log file, and streams the bytes back."
)
code_block("msg = consumer.poll(timeout=1.0)   # ask broker: 'any messages since my last offset?'")
body(
    "Each time the consumer reads a message it advances its local offset counter. "
    "When it commits, it writes that offset back to the broker (stored in the internal "
    "'__consumer_offsets' topic) so that if the consumer restarts it knows where to resume."
)

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 8 — TOPICS AND PARTITIONS
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Chapter 8 — Topics and Partitions")

h2("Topics")
body(
    "A topic is a named stream of messages. Think of it as a category. "
    "In this project the topic is called 'orders'. "
    "Producers write to a topic; consumers read from a topic."
)

h2("Partitions")
body(
    "A topic is split into one or more partitions. Each partition is an independent, "
    "ordered, append-only log. Partitions are the unit of parallelism in Kafka."
)
bullet("Within a single partition, messages are strictly ordered by offset.")
bullet("Across partitions of the same topic, there is NO global ordering guarantee.")
bullet("Each partition is stored as a separate set of segment files on disk.")

h2("Why have multiple partitions?")
bullet("Throughput: multiple producers can write to different partitions in parallel.")
bullet("Parallelism: multiple consumers (in the same group) can read different partitions in parallel — one consumer per partition max.")
bullet("Scalability: partitions can be spread across multiple brokers in a cluster.")

h2("How a message lands in a partition")
body("The partition is chosen by the producer based on the message key:")
bullet("Key provided: partition = hash(key) % num_partitions. Same key always goes to same partition.")
bullet("No key: round-robin across all partitions.")
body(
    "In your producer.py, the key is str(order_id). "
    "This means all events for the same order always land in the same partition, "
    "guaranteeing ordering for that order."
)
code_block('key=str(order["order_id"]).encode("utf-8")')

h2("Replication factor")
body(
    "In production each partition is replicated across N brokers (the replication factor). "
    "One broker holds the leader replica (handles all reads and writes); "
    "the others hold follower replicas (stay in sync, take over if the leader dies). "
    "In this project KAFKA_OFFSETS_TOPIC_REPLICATION_FACTOR is 1 because there is only one broker."
)

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 9 — CONSUMER FAILURE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Chapter 9 — What Happens if the Consumer Fails?")

h2("The data is not lost")
body(
    "Kafka is not a queue that deletes messages once consumed. "
    "Messages live on the broker's disk until the retention period expires "
    "(default: 7 days, or a configurable size limit). "
    "A consumer crash does not delete any messages."
)

h2("How resumption works — offsets")
body(
    "Every consumer group has its own set of committed offsets stored in the broker "
    "(in the internal '__consumer_offsets' topic). "
    "When your consumer crashes, the committed offset for 'order-processor' remains "
    "saved on the broker. When the consumer restarts it asks the broker: "
    "'What was my last committed offset for the orders topic?' and resumes from there."
)

h2("auto.offset.reset explained")
body("This setting only applies when a consumer group has NO stored offset yet (first run).")
bullet("earliest — start from the very first message in the topic (offset 0).")
bullet("latest — start from messages that arrive after the consumer starts.")
body("Once a group has committed at least one offset, this setting is ignored on restart.")

h2("At-least-once vs at-most-once delivery")
body(
    "If the consumer processes a message but crashes before committing the offset, "
    "it will re-read and re-process that message on restart. "
    "This is called at-least-once delivery — messages may be processed more than once. "
    "Your business logic should handle duplicates (idempotency) — for example, "
    "using the order_id as a primary key in Postgres so a duplicate insert is a no-op."
)

h2("Consumer group rebalance")
body(
    "When a consumer in a group dies (or a new one joins), Kafka triggers a rebalance: "
    "it redistributes the topic's partitions among the surviving consumers. "
    "This is why consumer.py calls consumer.close() in the finally block — "
    "a clean close triggers an immediate rebalance instead of waiting for a timeout "
    "(the session timeout is typically 10–30 seconds)."
)

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 10 — INCREASING PARTITIONS
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Chapter 10 — Increasing Partitions and Why Docker Commands Are Needed")

h2("Why you cannot just edit a config file")
body(
    "Partition count is a live property of a topic stored inside Kafka's metadata "
    "(managed by Zookeeper). It is not a static config file on disk. "
    "To change it you have to talk to the running broker using the Kafka admin API "
    "or the Kafka CLI tools that ship inside the broker container."
)

h2("How to increase partition count")
code_block(
    "docker exec -it kafka-kafka-1 \\\n"
    "  kafka-topics --bootstrap-server localhost:9092 \\\n"
    "  --alter --topic orders --partitions 6\n\n"
    "# docker exec -it kafka-kafka-1  — open a shell inside the running Kafka container\n"
    "# kafka-topics                   — the Kafka CLI tool for topic administration\n"
    "# --bootstrap-server localhost:9092  — connect to the broker (localhost inside the container)\n"
    "# --alter                        — modify an existing topic\n"
    "# --topic orders                 — the topic to change\n"
    "# --partitions 6                 — the new partition count (must be higher than current)"
)

h2("Why you cannot decrease partitions")
body(
    "Messages in a partition are ordered by offset. If you removed a partition, "
    "you would have to decide where to move its messages — breaking offset ordering. "
    "Kafka does not allow it. If you need fewer partitions, create a new topic."
)

h2("Other useful kafka-topics commands")
code_block(
    "# List all topics\n"
    "docker exec -it kafka-kafka-1 kafka-topics --bootstrap-server localhost:9092 --list\n\n"
    "# Describe a topic (partitions, replication, leader)\n"
    "docker exec -it kafka-kafka-1 kafka-topics --bootstrap-server localhost:9092 --describe --topic orders\n\n"
    "# Create a topic manually\n"
    "docker exec -it kafka-kafka-1 kafka-topics --bootstrap-server localhost:9092 \\\n"
    "  --create --topic payments --partitions 3 --replication-factor 1\n\n"
    "# Delete a topic\n"
    "docker exec -it kafka-kafka-1 kafka-topics --bootstrap-server localhost:9092 --delete --topic orders\n\n"
    "# Check consumer group lag\n"
    "docker exec -it kafka-kafka-1 kafka-consumer-groups --bootstrap-server localhost:9092 \\\n"
    "  --group order-processor --describe\n\n"
    "# Read messages directly from terminal (great for debugging)\n"
    "docker exec -it kafka-kafka-1 kafka-console-consumer \\\n"
    "  --bootstrap-server localhost:9092 --topic orders --from-beginning"
)

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 11 — WHERE ALL CONFIGS LIVE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Chapter 11 — Where Do All These Configs Live?")

h2("docker-compose.yml — infrastructure config")
body(
    "Single source of truth for your local infrastructure. "
    "Every KAFKA_* environment variable is translated by the Confluent Docker image "
    "into a line in Kafka's broker config file at startup."
)
code_block(
    "KAFKA_BROKER_ID: 1                        # unique ID for this broker in the cluster\n"
    "KAFKA_ZOOKEEPER_CONNECT: zookeeper:2181   # where to find Zookeeper\n"
    "KAFKA_ADVERTISED_LISTENERS: PLAINTEXT://localhost:9092\n"
    "KAFKA_LISTENERS: PLAINTEXT://0.0.0.0:9092\n"
    "KAFKA_AUTO_CREATE_TOPICS_ENABLE: 'true'\n"
    "KAFKA_OFFSETS_TOPIC_REPLICATION_FACTOR: 1"
)

h2(".env — application config")
body("Your Python scripts load this file via python-dotenv:")
code_block(
    "KAFKA_BOOTSTRAP_SERVERS=localhost:9092\n"
    "POSTGRES_HOST=localhost\n"
    "POSTGRES_PORT=5432\n"
    "POSTGRES_DB=your_db_name\n"
    "POSTGRES_USER=your_user\n"
    "POSTGRES_PASSWORD=your_password"
)
note("Never commit real credentials to git. Add .env to .gitignore.")

h2("Where Kafka's config file lives inside the container")
code_block("/etc/kafka/kafka.properties   (inside the container)")
body("To inspect it:")
code_block("docker exec -it kafka-kafka-1 cat /etc/kafka/kafka.properties")

h2("Where consumer group offsets are stored")
body("Committed offsets live on the broker in the internal topic '__consumer_offsets':")
code_block(
    "docker exec -it kafka-kafka-1 \\\n"
    "  kafka-consumer-groups --bootstrap-server localhost:9092 \\\n"
    "  --group order-processor --describe\n\n"
    "# Prints: partition, current offset, log-end offset, lag"
)

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 12 — PROJECT FILES ANNOTATED
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Chapter 12 — Your Project Files Explained")

h2("producer.py — annotated walkthrough")
code_block(
    "producer = Producer({\n"
    "    'bootstrap.servers': 'localhost:9092',  # only required setting\n"
    "})"
)
body(
    "bootstrap.servers is just the initial contact point. "
    "The producer fetches full cluster metadata from this broker "
    "and then connects directly to partition leaders."
)
code_block(
    "producer.produce(\n"
    "    topic=TOPIC,\n"
    "    value=json.dumps(order).encode('utf-8'),  # messages are raw bytes\n"
    "    key=str(order['order_id']).encode('utf-8'),  # same key -> same partition\n"
    "    callback=delivery_report,\n"
    ")"
)
body(
    "The message is not sent immediately. It goes into an internal buffer. "
    "producer.poll(0) fires delivery callbacks; "
    "producer.flush() blocks until all buffered messages are acknowledged by the broker."
)

h2("consumer.py — annotated walkthrough")
code_block(
    "consumer = Consumer({\n"
    "    'bootstrap.servers': 'localhost:9092',\n"
    "    'group.id': 'order-processor',   # identifies this consumer's group\n"
    "    'auto.offset.reset': 'earliest', # first run: start from beginning\n"
    "})\n\n"
    "consumer.subscribe(['orders'])       # Kafka assigns partitions to this consumer\n"
    "msg = consumer.poll(timeout=1.0)     # block up to 1s waiting for a message"
)
body(
    "auto.commit is on by default. The consumer commits its offset back to the broker "
    "periodically (every 5 seconds). You can disable this and commit manually for "
    "exactly-once semantics."
)

h2("inspect_topics.py — annotated walkthrough")
code_block(
    "admin = AdminClient({'bootstrap.servers': BOOTSTRAP})\n"
    "metadata = admin.list_topics(timeout=10)  # fetches all topic/partition metadata\n\n"
    "low, high = consumer.get_watermark_offsets(tp, timeout=5)\n"
    "# low  = earliest available offset (may be > 0 if old segments were deleted)\n"
    "# high = next offset to be written (= total messages if no segment deletions)\n"
    "lag = high - committed_offset"
)
body(
    "Lag is the number of messages the consumer group has not yet read. "
    "If lag is growing, the consumer is falling behind the producer."
)

# ═════════════════════════════════════════════════════════════════════════════
# CHAPTER 13 — QUICK REFERENCE
# ═════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Chapter 13 — Quick Reference")

h2("Docker standalone commands")
code_block(
    "docker pull <image>          # download an image from Docker Hub\n"
    "docker images                # list locally downloaded images\n"
    "docker rmi <image>           # delete a local image\n"
    "docker run -d -p H:C <image> # create and start a container\n"
    "docker ps                    # list running containers\n"
    "docker ps -a                 # list all containers (including stopped)\n"
    "docker stop <name>           # gracefully stop a container\n"
    "docker start <name>          # restart a stopped container\n"
    "docker rm <name>             # delete a stopped container\n"
    "docker rm -f <name>          # force delete (stop + remove)\n"
    "docker exec -it <name> bash  # open a shell inside a running container\n"
    "docker logs -f <name>        # follow a container's logs\n"
    "docker inspect <name>        # detailed JSON info about a container\n"
    "docker volume ls             # list volumes\n"
    "docker volume rm <vol>       # delete a volume\n"
    "docker system prune          # clean up stopped containers + dangling images\n"
    "docker system prune -a       # also clean unused images"
)

h2("Docker Compose commands")
code_block(
    "docker compose up -d         # start all services in background\n"
    "docker compose down          # stop and remove containers\n"
    "docker compose down -v       # also delete volumes (wipes all data)\n"
    "docker compose stop          # stop containers but keep them\n"
    "docker compose start         # restart stopped containers\n"
    "docker compose restart kafka # restart one service\n"
    "docker compose ps            # status of all services\n"
    "docker compose logs -f kafka # live logs for kafka service\n"
    "docker compose exec kafka bash  # shell inside a running service\n"
    "docker compose pull          # update all images"
)

h2("Kafka CLI (run inside the container)")
code_block(
    "# List topics\n"
    "docker exec -it kafka-kafka-1 kafka-topics --bootstrap-server localhost:9092 --list\n\n"
    "# Describe topic\n"
    "docker exec -it kafka-kafka-1 kafka-topics --bootstrap-server localhost:9092 --describe --topic orders\n\n"
    "# Increase partitions\n"
    "docker exec -it kafka-kafka-1 kafka-topics --bootstrap-server localhost:9092 --alter --topic orders --partitions 4\n\n"
    "# Check consumer group lag\n"
    "docker exec -it kafka-kafka-1 kafka-consumer-groups --bootstrap-server localhost:9092 --group order-processor --describe\n\n"
    "# Read messages from beginning (debug)\n"
    "docker exec -it kafka-kafka-1 kafka-console-consumer --bootstrap-server localhost:9092 --topic orders --from-beginning"
)

h2("Run your Python scripts")
code_block(
    "pip3 install -r requirements.txt\n"
    "python3 01_basics/producer.py        # send 59 orders\n"
    "python3 01_basics/consumer.py        # start consuming (Ctrl+C to stop)\n"
    "python3 01_basics/inspect_topics.py  # check lag"
)

h2("Key concepts cheat sheet")
bullet("Image — read-only blueprint for a container (like a class)")
bullet("Container — a running instance of an image (like an object)")
bullet("Volume — persistent storage that survives container restarts")
bullet("Port mapping HOST:CONTAINER — bridges container ports to your laptop")
bullet("docker compose up -d — start everything defined in docker-compose.yml")
bullet("Topic — named stream of messages in Kafka")
bullet("Partition — ordered sub-log of a topic; unit of parallelism")
bullet("Offset — permanent position of a message in a partition")
bullet("Consumer Group — set of consumers sharing work; each partition read by exactly one")
bullet("Lag — messages produced but not yet consumed by a group")
bullet("Retention — how long Kafka keeps messages (default 7 days)")
bullet("Broker — the Kafka server process that stores and serves messages")
bullet("Bootstrap servers — initial contact point; broker returns full cluster metadata")
bullet("Replication factor — how many copies of each partition across brokers")
bullet("Leader — the one broker replica that handles all reads/writes for a partition")

# ── Save ──────────────────────────────────────────────────────────────────────
out = "/Users/rafidbinsadeque/Desktop/kafka/kafka_explained.docx"
doc.save(out)
print(f"Saved: {out}")
